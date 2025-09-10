# app_horizontal_tank.py
# Dimensionare rezervor orizontal (cilindru + capace elipsoidale 2:1)
# Organizare intrări pe grupuri: Capacitate & Geometrie / Material & Mecanică / Saddle-uri & Montaj / Calibrare & Grafic

import math
import io
import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
from openpyxl import Workbook

st.set_page_config(page_title="Rezervor orizontal - dimensionare", layout="wide")

# =========================== Funcții auxiliare ===========================
def head_volume_21_ellipsoidal(R: float) -> float:
    """Volum capac 2:1 elipsoidal: V_head = (1/3)*π*R^3"""
    return (math.pi / 3.0) * (R ** 3)

def level_segment_area(R: float, h: float) -> float:
    """Aria segmentului circular pentru înălțime h în cilindru (0..2R)."""
    if h <= 0:
        return 0.0
    if h >= 2 * R:
        return math.pi * R ** 2
    return R**2 * math.acos((R - h) / R) - (R - h) * math.sqrt(max(0.0, 2*R*h - h**2))

def export_excel(inputs, geom, df_lv, mech, saddles) -> io.BytesIO:
    wb = Workbook()

    # Sheet 1: Date de intrare (rezumat grupat)
    ws = wb.active
    ws.title = "Setări & Rezumat"
    ws["A1"] = "SETĂRI INTRARE (grupate)"
    r = 3
    for sect, dct in inputs.items():
        ws.cell(r, 1).value = f"[{sect}]"
        r += 1
        for k, v in dct.items():
            ws.cell(r, 1).value = k
            ws.cell(r, 2).value = float(v) if isinstance(v, (int, float)) else v
            r += 1
        r += 1

    # Sheet 2: Geometrie
    ws2 = wb.create_sheet("Geometrie")
    r = 1
    for k, v in geom.items():
        ws2.cell(r, 1).value = k
        ws2.cell(r, 2).value = float(v) if isinstance(v, (int, float)) else v
        r += 1

    # Sheet 3: Nivel–Volum
    ws3 = wb.create_sheet("Nivel–Volum")
    ws3.append(list(df_lv.columns))
    for row in df_lv.itertuples(index=False):
        ws3.append(list(row))

    # Sheet 4: Verificare mecanică
    ws4 = wb.create_sheet("Verificare mecanică")
    r = 1
    for k, v in mech.items():
        ws4.cell(r, 1).value = k
        ws4.cell(r, 2).value = float(v) if isinstance(v, (int, float)) else v
        r += 1

    # Sheet 5: Saddle-uri
    ws5 = wb.create_sheet("Saddle-uri")
    r = 1
    for k, v in saddles.items():
        ws5.cell(r, 1).value = k
        ws5.cell(r, 2).value = float(v) if isinstance(v, (int, float)) else v
        r += 1

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf

# =========================== Sidebar – GRUPURI ===========================
with st.sidebar:
    st.title("Date de intrare")

    # A) Capacitate & Geometrie
    with st.expander("A) Capacitate & Geometrie", expanded=True):
        V_work = st.number_input("Volum util necesar V_util [m³]", value=100.0, min_value=0.0, step=1.0)
        f_fill = st.slider("Fracție de umplere f_umplere [-]", 0.50, 1.00, 0.90, 0.01, format="%.2f")
        V_total_target = V_work / f_fill if f_fill > 0 else 0.0
        st.caption(f"Volum total de proiectare **V_total = {V_total_target:.3f} m³**")

        mode = st.selectbox("Mod de dimensionare", ("Fixez L/D și calculez volumul rezultat", "Fixez D și calculez lungimea pentru V_total"))
        colA, colB = st.columns([3, 2])
        with colA:
            L_over_D = st.slider("Raport L/D (recomandat 2–5)", 1.50, 6.00, 3.00, step=0.01, format="%.2f")
        with colB:
            L_over_D = st.number_input("L/D (tastat fin)", min_value=1.50, max_value=6.00, value=float(L_over_D), step=0.01, format="%.2f")

        D = st.number_input("Diametru D [m]", value=3.2, min_value=0.1, step=0.05, help="D poate fi limitat de transport.")
        D_max = st.number_input("Diametru maxim admis D_max [m] (opțional)", value=0.0, min_value=0.0, step=0.1, help="Lasă 0 dacă nu ai constrângere.")
        L_max = st.number_input("Lungime totală maximă L_max [m] (opțional)", value=0.0, min_value=0.0, step=0.5, help="Lasă 0 dacă nu ai constrângere.")

    # B) Material & Mecanică
    with st.expander("B) Material & Mecanică", expanded=False):
        rho = st.number_input("Densitate lichid ρ [kg/m³]", value=1000.0, min_value=1.0)
        sigma_allow = st.number_input("Tensiune admisă material σ_adm [MPa]", value=120.0, min_value=1.0)
        SF = st.number_input("Factor siguranță SF [-]", value=1.5, min_value=1.0)
        t_min = st.number_input("Grosime minimă de fabricație t_min [mm]", value=6.0, min_value=0.0, step=0.5)
        CA = st.number_input("Adaos de coroziune CA [mm]", value=1.0, min_value=0.0, step=0.5)
        t_nom = st.number_input("Grosime nominală aleasă t_nom [mm]", value=8.0, min_value=0.0, step=0.5)

    # C) Saddle-uri & Montaj
    with st.expander("C) Saddle-uri & Montaj", expanded=False):
        a = st.number_input("Distanță de la tangentă la centrul saddle a [m]", value=0.2, min_value=0.0, step=0.05)
        S_user = st.number_input("S ales (span centru–centru) [m] (opțional)", value=0.0, min_value=0.0, step=0.05, help="Poți lăsa 0 și vei vedea intervalul recomandat în tab-ul Saddle-uri.")

    # D) Calibrare & Grafic
    with st.expander("D) Calibrare & Grafic", expanded=False):
        n_points = st.slider("Număr puncte pe curba nivel–volum", 20, 300, 120, 5)

# =========================== Calcul Geometrie în funcție de MOD ===========================
R = D / 2.0
h_head = D / 4.0                 # 2:1 elipsoidal
V_head = head_volume_21_ellipsoidal(R)

if mode.startswith("Fixez D"):
    # Calculez L_cil pentru a obține volumul țintă
    V_cyl_needed = max(0.0, V_total_target - 2.0 * V_head)
    L_cyl = V_cyl_needed / (math.pi * R**2) if R > 0 else 0.0
    L_total = L_cyl + 2.0 * h_head
    L_over_D_calc = L_total / D if D > 0 else 0.0
    V_total_calc = V_cyl_needed + 2.0 * V_head
    dV = V_total_calc - V_total_target
else:
    # Fixez L/D -> obțin L_total, apoi volumul rezultat (poate fi diferit de țintă)
    L_total = L_over_D * D
    L_cyl = max(0.0, L_total - 2.0 * h_head)
    V_cyl = math.pi * R**2 * L_cyl
    V_total_calc = V_cyl + 2.0 * V_head
    dV = V_total_calc - V_total_target
    L_over_D_calc = L_total / D if D > 0 else 0.0

# check-uri geometrice
warn_D = (D_max > 0) and (D > D_max)
warn_L = (L_max > 0) and (L_total > L_max)

# =========================== Tabs UI ===========================
tab_geom, tab_lv, tab_mech, tab_sad, tab_sketch = st.tabs(
    ["Geometrie", "Nivel–Volum", "Verificare mecanică", "Saddle-uri", "Schiță"]
)

# -------- Geometrie --------
with tab_geom:
    st.subheader("Geometrie & Capacitate")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Rază R [m]", f"{R:.3f}")
    c2.metric("Adâncime capac h_cap [m]", f"{h_head:.3f}")
    c3.metric("L_cil [m]", f"{L_cyl:.3f}")
    c4.metric("L_total [m]", f"{L_total:.3f}")

    c1.metric("V_cap (1) [m³]", f"{V_head:.3f}")
    c2.metric("V_total calculat [m³]", f"{V_total_calc:.3f}")
    c3.metric("ΔV = V_calc - V_țintă [m³]", f"{dV:.3f}")
    c4.metric("L/D rezultat [-]", f"{L_over_D_calc:.2f}")

    if warn_D:
        st.error(f"Depășește D_max: D = {D:.2f} m > D_max = {D_max:.2f} m")
    if warn_L:
        st.error(f"Depășește L_max: L_total = {L_total:.2f} m > L_max = {L_max:.2f} m")

    st.info("Ajustează L/D sau D în funcție de modul selectat astfel încât ΔV ≈ 0 și să respecți D_max / L_max.")

# -------- Nivel–Volum --------
with tab_lv:
    st.subheader("Curba NIVEL–VOLUM (cilindru + capace ca termen constant)")
    h_vals = np.linspace(0.0, D, n_points)
    A_vals = np.array([level_segment_area(R, h) for h in h_vals])
    V_cyl_h = A_vals * L_cyl
    V_total_h = V_cyl_h + 2.0 * V_head
    df_lv = pd.DataFrame({"h [m]": h_vals, "Volum_cil [m³]": V_cyl_h, "Volum_total [m³]": V_total_h})
    st.dataframe(df_lv.head(12))

    fig, ax = plt.subplots()
    ax.plot(df_lv["h [m]"], df_lv["Volum_total [m³]"])
    ax.set_xlabel("Înălțime lichid h [m]")
    ax.set_ylabel("Volum total [m³]")
    ax.set_title("Curba NIVEL–VOLUM (orizontal)")
    st.pyplot(fig)

# -------- Verificare mecanică --------
with tab_mech:
    st.subheader("Verificare simplificată – efort cerc hidrostatic (rezervor atmosferic)")
    sigma_used = sigma_allow / SF
    h_design = D  # coloană reprezentativă
    t_req_mm = (rho * 9.80665 * h_design * R) / (sigma_used * 1000.0)
    t_disp_mm = t_nom - CA
    need = max(t_req_mm, t_min)
    ok = t_disp_mm >= need

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("σ utilizată [MPa]", f"{sigma_used:.2f}")
    c2.metric("t_req [mm]", f"{t_req_mm:.2f}")
    c3.metric("t_disp = t_nom - CA [mm]", f"{t_disp_mm:.2f}")
    c4.metric("Necesar ≥ max(t_req, t_min)", f"{need:.2f}")

    if ok:
        st.success("OK – grosimea este suficientă")
    else:
        st.error("Crește grosimea – t_disp < max(t_req, t_min)")

# -------- Saddle-uri --------
with tab_sad:
    st.subheader("Saddle-uri (simplificat)")
    S_rec_low = 0.4 * L_cyl
    S_rec_high = 0.5 * L_cyl
    span = L_total - 2.0 * a

    S_chosen = S_user if S_user > 0 else S_rec_low
    c1, c2, c3 = st.columns(3)
    c1.metric("Interval recomandat S [m]", f"{S_rec_low:.3f} – {S_rec_high:.3f}")
    c2.metric("Deschidere totală L_total - 2a [m]", f"{span:.3f}")
    c3.metric("S ales [m]", f"{S_chosen:.3f}")

    if S_chosen <= span:
        st.info("Check: S ≤ deschidere → OK")
    else:
        st.error("Check: S > deschidere → Ajustează distanța / offsetul a")

# -------- Schiță --------
with tab_sketch:
    st.subheader("Schiță rezervor orizontal – capace elipsoidale 2:1 (la scară)")
    from matplotlib.patches import Rectangle, Ellipse

    # Helperi pentru cote (text deasupra)
    def cota_oriz(ax, x1, x2, y, label, dy_text=0.08, color="black"):
        ax.annotate("", xy=(x1, y), xytext=(x2, y),
                    arrowprops=dict(arrowstyle="<->", lw=1.4, color=color, shrinkA=4, shrinkB=4))
        ax.text((x1+x2)/2, y + dy_text * D, label, ha="center", va="bottom", fontsize=11, color=color,
                bbox=dict(facecolor="white", edgecolor="none", pad=1.5))

    def cota_vert(ax, x, y1, y2, label, dx_text=0.06, color="black"):
        ax.annotate("", xy=(x, y1), xytext=(x, y2),
                    arrowprops=dict(arrowstyle="<->", lw=1.4, color=color, shrinkA=4, shrinkB=4))
        ax.text(x + dx_text * D, (y1+y2)/2, label, ha="left", va="center", rotation=90, fontsize=11, color=color,
                bbox=dict(facecolor="white", edgecolor="none", pad=1.5))

    # Geometrie pentru desen
    Rplot = R; Dplot = D; Lc = L_cyl; Ltot = L_total; hcap = h_head
    x_tan_L = hcap; x_tan_R = hcap + Lc
    x_s1 = x_tan_L - a
    S_chosen = S_user if S_user > 0 else 0.4 * L_cyl
    x_s2 = x_tan_L + S_chosen
    span_av = Ltot - 2 * a

    fig, ax = plt.subplots(figsize=(11, 3.6))

    # Cilindru + capace 2:1 (elipse)
    ax.add_patch(Rectangle((x_tan_L, -Rplot), Lc, 2*Rplot, fill=False, lw=2.5, joinstyle="round"))
    ax.add_patch(Ellipse((x_tan_L, 0), width=2*hcap, height=2*Rplot, angle=0, fill=False, lw=2.5))
    ax.add_patch(Ellipse((x_tan_R, 0), width=2*hcap, height=2*Rplot, angle=0, fill=False, lw=2.5))
    ax.plot([x_tan_L, x_tan_L], [-Rplot, Rplot], color="black", lw=2.5)
    ax.plot([x_tan_R, x_tan_R], [-Rplot, Rplot], color="black", lw=2.5)
    ax.axhline(0, color="0.5", lw=0.8, ls="--", alpha=0.5)

    # Saddle-uri
    pad_h = 0.15 * Rplot
    ax.add_patch(Rectangle((x_s1 - 0.05 * Dplot, -Rplot - pad_h), 0.10 * Dplot, pad_h, color="#ff8c00"))
    ax.add_patch(Rectangle((x_s2 - 0.05 * Dplot, -Rplot - pad_h), 0.10 * Dplot, pad_h, color="#2e7d32"))
    ax.text(x_s1, -Rplot - pad_h - 0.22 * Dplot, "Saddle 1", ha="center", va="top", fontsize=10)
    ax.text(x_s2, -Rplot - pad_h - 0.22 * Dplot, "Saddle 2", ha="center", va="top", fontsize=10)

    # Cote
    cota_vert(ax, x_tan_L + Lc/2, -Rplot, Rplot, f"D = {D:.2f} m")
    cota_oriz(ax, x_tan_L, x_tan_R, 1.28*Rplot, f"L_cil = {Lc:.2f} m", dy_text=0.06)
    cota_oriz(ax, 0, Ltot, -1.45*Rplot, f"L_total = {Ltot:.2f} m", dy_text=0.06)
    cota_oriz(ax, x_s1, x_tan_L, -1.05*Rplot, f"a = {a:.2f} m", dy_text=0.06)
    cota_oriz(ax, x_s1, x_s2, -1.25*Rplot, f"S ales = {S_chosen:.2f} m", dy_text=0.06, color="tab:blue")
    cota_oriz(ax, x_tan_L - a, x_tan_R + a, -1.65*Rplot, f"Deschidere = {span_av:.2f} m", dy_text=0.06, color="tab:green")

    ax.set_aspect("equal", adjustable="box")
    ax.set_xlim(-0.35 * Dplot, Ltot + 0.35 * Dplot)
    ax.set_ylim(-2.0 * Rplot, 1.75 * Rplot)
    ax.axis("off")
    st.pyplot(fig)

    # Export PNG
    buf_img = io.BytesIO()
    fig.savefig(buf_img, format="png", dpi=200, bbox_inches="tight")
    buf_img.seek(0)
    st.download_button("Descarcă schița (PNG)", data=buf_img, file_name="schita_rezervor_orizontal.png", mime="image/png")

# =========================== Export Word (DOCX) – Concluzii ===========================
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def make_conclusions_docx(geom, mech, saddles) -> bytes:
    """
    Generează un DOCX cu concluzii concise: Geometrie, Verificare mecanică, Saddle-uri.
    Parametrii așteptați:
      geom: dict cu chei R [m], h_cap [m], L_cil [m], L_total [m], V_head (1) [m³], V_total calc [m³], ΔV [m³], L/D rezultat [-]
      mech: dict cu chei σ utilizată [MPa], h design [m], t_req [mm], t_min [mm], CA [mm], t_nom [mm], t_disp [mm]
      saddles: dict cu chei S_rec_min [m], S_rec_max [m], Deschidere [m]
    """
    d = Document()

    # Titlu
    title = d.add_paragraph("Raport scurt – Rezervor orizontal (cilindru + capace 2:1)")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]; run.bold = True; run.font.size = Pt(14)

    # Geometrie
    d.add_heading("1) Geometrie & Capacitate", level=2)
    p = d.add_paragraph()
    p.add_run(f"Diametru D: ").bold = True; p.add_run(f"{2*geom['R [m]']:.2f} m\n")
    p.add_run(f"Lungime cilindru L_cil: ").bold = True; p.add_run(f"{geom['L_cil [m]']:.2f} m\n")
    p.add_run(f"Lungime totală L_total: ").bold = True; p.add_run(f"{geom['L_total [m]']:.2f} m\n")
    p.add_run(f"Volum total calculat: ").bold = True; p.add_run(f"{geom['V_total calc [m³]']:.2f} m³\n")
    p.add_run(f"Abatere ΔV (față de țintă): ").bold = True; 
    dv = geom["ΔV [m³]"]; p.add_run(f"{dv:+.2f} m³  ").bold = False
    # verdict ΔV
    if abs(dv) <= 0.5:
        p.add_run("(OK – în toleranță)").italic = True
    else:
        p.add_run("(Ajustează D/L/D pentru ΔV≈0)").italic = True

    # Verificare mecanică
    d.add_heading("2) Verificare mecanică (efort cerc – hidrostatic)", level=2)
    ok_mech = (mech["t_disp [mm]"] >= max(mech["t_req [mm]"], mech["t_min [mm]"]))
    tbl = d.add_table(rows=4, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "σ utilizată [MPa]"; hdr[1].text = "t_req [mm]"
    hdr[2].text = "t_min [mm]"; hdr[3].text = "t_disp = t_nom − CA [mm]"
    row = tbl.rows[1].cells
    row[0].text = f"{mech['σ utilizată [MPa]']:.2f}"
    row[1].text = f"{mech['t_req [mm]']:.2f}"
    row[2].text = f"{mech['t_min [mm]']:.2f}"
    row[3].text = f"{mech['t_disp [mm]']:.2f}"

    verdict = d.add_paragraph()
    verdict_run = verdict.add_run("Verdict: ")
    verdict_run.bold = True
    verdict.add_run("OK – grosimea este suficientă" if ok_mech else "NU – crește grosimea (t_disp < max(t_req, t_min))").bold = True

    # Saddle-uri
    d.add_heading("3) Saddle-uri (dispunere)", level=2)
    p2 = d.add_paragraph()
    p2.add_run("Distanță recomandată S: ").bold = True
    p2.add_run(f"{saddles['S_rec_min [m]']:.2f} – {saddles['S_rec_max [m]']:.2f} m\n")
    p2.add_run("Deschidere disponibilă L_total − 2a: ").bold = True
    p2.add_run(f"{saddles['Deschidere [m]']:.2f} m\n")
    p2.add_run("Notă: alege S în intervalul recomandat și verifică S ≤ deschidere.").italic = True

    # Footer scurt
    d.add_paragraph().add_run("Acest raport sumarizează rezultatele principale; pentru detalii vedeți aplicația.").italic = True

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf.read()

# Pregătim dicționarele deja create mai sus (geom_dict, mech_dict, saddles_dict, df_lv existent)
docx_bytes = make_conclusions_docx(geom_dict, mech_dict, saddles_dict)
st.download_button("Descarcă raport (Word)", data=docx_bytes,
                   file_name="Raport_rezervor_orizontal.docx",
                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
